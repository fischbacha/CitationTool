from __future__ import annotations

import json
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class BuildOutputs:
    out_dir: Path
    placeholder_docx: Path
    active_docx: Path
    ris: Path
    csl_json: Path
    claim_report: Path


def load_project(spec_path: Path) -> dict[str, Any]:
    project = json.loads(spec_path.read_text(encoding="utf-8"))
    base_dir = spec_path.parent
    source = project.get("reference_source")
    if source:
        source_path = (base_dir / source).resolve()
        if not source_path.exists():
            source_path = (ROOT / source).resolve()
        project["references"] = json.loads(source_path.read_text(encoding="utf-8"))
    if not project.get("references"):
        raise ValueError("Project spec needs either `references` or `reference_source`.")
    if not project.get("paragraphs"):
        raise ValueError("Project spec needs `paragraphs`.")
    project.setdefault("slug", "citationtool_project")
    project.setdefault("title", project["slug"].replace("_", " ").title())
    project.setdefault("subtitle", "Generated Zotero-active Word draft")
    project.setdefault("reference_heading", "References")
    project.setdefault("body_heading", "Draft Introduction")
    return project


def project_out_dir(project: dict[str, Any], override: Path | None = None) -> Path:
    if override:
        return override
    if project.get("output_dir"):
        return (ROOT / project["output_dir"]).resolve()
    return ROOT / "artifacts" / project["slug"]


def output_paths(project: dict[str, Any], out_dir: Path) -> BuildOutputs:
    slug = project["slug"]
    return BuildOutputs(
        out_dir=out_dir,
        placeholder_docx=out_dir / f"{slug}_zotero_placeholders.docx",
        active_docx=out_dir / f"{slug}_zotero_active.docx",
        ris=out_dir / f"{slug}_references.ris",
        csl_json=out_dir / f"{slug}_references.csl.json",
        claim_report=out_dir / "claim_support_report.md",
    )


def set_font(run, name: str = "Calibri", size: int | float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name: str, size: int | float, color: str | None = None, bold: bool | None = None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 20, "0B2545", True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 11, "555555")
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, "2E74B5", True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, "2E74B5", True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)


def configure_section(doc: Document):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    return section


def add_title_block(doc: Document, project: dict[str, Any], active: bool):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run(project["title"])
    set_font(title_run, size=20, bold=True)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_text = project.get("active_subtitle" if active else "placeholder_subtitle") or project["subtitle"]
    subtitle_run = subtitle.add_run(subtitle_text)
    set_font(subtitle_run, size=11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")


def references_by_key(project: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {ref["id"]: ref for ref in project["references"]}


def citation_numbers(project: dict[str, Any]) -> dict[str, int]:
    return {ref["id"]: index for index, ref in enumerate(project["references"], 1)}


def author_to_ris(author: dict[str, Any]) -> str:
    family = author.get("family", "")
    given = author.get("given", "")
    return f"{family}, {given}".strip().strip(",")


def year_from_ref(ref: dict[str, Any]) -> str:
    parts = ref.get("issued", {}).get("date-parts", [])
    if parts and parts[0]:
        return str(parts[0][0])
    return str(ref.get("year", ""))


def page_bounds(ref: dict[str, Any]) -> tuple[str, str]:
    page = ref.get("page", "")
    if "-" in page:
        start, end = page.split("-", 1)
        return start, end
    return page, ""


def placeholder_for(ref: dict[str, Any], project: dict[str, Any]) -> str:
    placeholders = project.get("placeholders", {})
    if ref["id"] in placeholders:
        return placeholders[ref["id"]]
    author = ref.get("author", [{}])[0].get("family", ref["id"])
    return f"[ZOTERO: {author} {year_from_ref(ref)}]"


def visible_bibliography_entry(ref: dict[str, Any]) -> str:
    authors = ref.get("author", [])
    if not authors:
        author_label = ref.get("id", "Unknown")
    elif len(authors) > 2:
        author_label = f"{authors[0].get('family', 'Unknown')} et al."
    elif len(authors) == 2:
        author_label = f"{authors[0].get('family', 'Unknown')} and {authors[1].get('family', 'Unknown')}"
    else:
        author_label = authors[0].get("family", "Unknown")

    journal = ref.get("container-title-short") or ref.get("container-title", "")
    volume = ref.get("volume", "")
    issue = ref.get("issue", "")
    pages = ref.get("page", "")
    doi = ref.get("DOI", "")
    pmid = ref.get("PMID", "")
    volume_issue = f"{volume}({issue})" if issue else volume
    identifiers = []
    if doi:
        identifiers.append(f"doi:{doi}")
    if pmid:
        identifiers.append(f"PMID:{pmid}")
    tail = "; ".join(identifiers)
    if tail:
        tail = f" {tail}."
    return f"{author_label} {year_from_ref(ref)}. {ref.get('title', '')}. {journal} {volume_issue}:{pages}.{tail}"


def append_text_element(run_element, text: str):
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)


def append_field_run(paragraph, instruction: str, result_text: str):
    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    paragraph._p.append(begin)

    for idx in range(0, len(instruction), 900):
        instr_run = OxmlElement("w:r")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction[idx : idx + 900]
        instr_run.append(instr_text)
        paragraph._p.append(instr_run)

    separate = OxmlElement("w:r")
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate.append(fld_separate)
    paragraph._p.append(separate)

    for line_index, line in enumerate(result_text.split("\n")):
        result_run = OxmlElement("w:r")
        if line_index:
            br = OxmlElement("w:br")
            result_run.append(br)
        append_text_element(result_run, line)
        paragraph._p.append(result_run)

    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._p.append(end)


def zotero_like_uris(ref: dict[str, Any]) -> list[str]:
    key = ref["id"].replace("-", "").upper()[:8]
    return [f"http://zotero.org/users/local/citationtool/items/{key}"]


def zotero_citation_json(ref: dict[str, Any], number: int) -> str:
    visible = f"[{number}]"
    doi_or_id = ref.get("DOI") or ref["id"]
    payload = {
        "citationID": f"citationtool-{uuid.uuid5(uuid.NAMESPACE_URL, doi_or_id).hex[:12]}",
        "citationItems": [
            {
                "id": ref["id"],
                "uris": zotero_like_uris(ref),
                "itemData": ref,
            }
        ],
        "properties": {
            "formattedCitation": visible,
            "plainCitation": visible,
            "noteIndex": 0,
        },
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return json.dumps(payload, separators=(",", ":"), ensure_ascii=True)


def append_zotero_citation(paragraph, ref: dict[str, Any], number: int):
    instruction = f" ADDIN ZOTERO_ITEM CSL_CITATION {zotero_citation_json(ref, number)}"
    append_field_run(paragraph, instruction, f"[{number}]")


def append_zotero_bibliography(paragraph, bibliography_text: str):
    code = json.dumps({"uncited": [], "omitted": [], "custom": []}, separators=(",", ":"))
    instruction = f" ADDIN ZOTERO_BIBL {code} CSL_BIBLIOGRAPHY"
    append_field_run(paragraph, instruction, bibliography_text)


def add_active_paragraph(doc: Document, chunks: list[dict[str, Any]], project: dict[str, Any]):
    refs = references_by_key(project)
    numbers = citation_numbers(project)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.25
    for chunk in chunks:
        if "cite" in chunk:
            ref = refs[chunk["cite"]]
            append_zotero_citation(paragraph, ref, numbers[ref["id"]])
        else:
            run = paragraph.add_run(chunk.get("text", ""))
            set_font(run, size=11)


def add_placeholder_paragraph(doc: Document, chunks: list[dict[str, Any]], project: dict[str, Any]):
    refs = references_by_key(project)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.25
    for chunk in chunks:
        if "cite" in chunk:
            run = paragraph.add_run(placeholder_for(refs[chunk["cite"]], project))
        else:
            run = paragraph.add_run(chunk.get("text", ""))
        set_font(run, size=11)


def make_placeholder_docx(project: dict[str, Any], path: Path):
    doc = Document()
    section = configure_section(doc)
    configure_styles(doc)
    add_title_block(doc, project, active=False)
    doc.add_heading(project["body_heading"], level=1)
    for paragraph_chunks in project["paragraphs"]:
        add_placeholder_paragraph(doc, paragraph_chunks, project)
    doc.add_heading(project["reference_heading"], level=1)
    hint = doc.add_paragraph()
    run = hint.add_run(
        "[Fallback draft: replace each visible placeholder with Zotero Add/Edit Citation, then insert bibliography.]"
    )
    set_font(run, size=10.5)
    run.italic = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Generated CitationTool placeholder draft")
    set_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor.from_string("555555")
    doc.save(path)


def make_active_docx(project: dict[str, Any], path: Path):
    doc = Document()
    section = configure_section(doc)
    configure_styles(doc)
    add_title_block(doc, project, active=True)
    doc.add_heading(project["body_heading"], level=1)
    for paragraph_chunks in project["paragraphs"]:
        add_active_paragraph(doc, paragraph_chunks, project)
    doc.add_heading(project["reference_heading"], level=1)
    bibliography_text = "\n".join(
        f"{index}. {visible_bibliography_entry(ref)}" for index, ref in enumerate(project["references"], 1)
    )
    bibliography = doc.add_paragraph()
    bibliography.paragraph_format.space_after = Pt(6)
    bibliography.paragraph_format.line_spacing = 1.15
    append_zotero_bibliography(bibliography, bibliography_text)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Generated CitationTool Zotero-active draft")
    set_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor.from_string("555555")
    doc.save(path)


def make_csl_json(project: dict[str, Any], path: Path):
    path.write_text(json.dumps(project["references"], indent=2) + "\n", encoding="utf-8")


def make_ris(project: dict[str, Any], path: Path):
    blocks: list[str] = []
    for ref in project["references"]:
        start_page, end_page = page_bounds(ref)
        lines = ["TY  - JOUR", f"ID  - {ref['id']}"]
        for author in ref.get("author", []):
            lines.append(f"AU  - {author_to_ris(author)}")
        lines.extend(
            [
                f"TI  - {ref.get('title', '')}",
                f"T2  - {ref.get('container-title', '')}",
                f"JO  - {ref.get('container-title-short', ref.get('container-title', ''))}",
                f"PY  - {year_from_ref(ref)}",
                f"VL  - {ref.get('volume', '')}",
                f"IS  - {ref.get('issue', '')}",
                f"SP  - {start_page}",
                f"EP  - {end_page}",
                f"DO  - {ref.get('DOI', '')}",
                f"UR  - {ref.get('URL', '')}",
            ]
        )
        if ref.get("PMID"):
            lines.append(f"AN  - PMID:{ref['PMID']}")
        lines.append("ER  -")
        blocks.append("\n".join(lines))
    path.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")


def make_claim_report(project: dict[str, Any], path: Path):
    refs = references_by_key(project)
    lines = [
        "# Claim Support Report",
        "",
        f"Generated for `{project['slug']}`.",
        "",
        "## References",
        "",
        "| Key | Verified identifiers | Reference |",
        "|---|---|---|",
    ]
    for ref in project["references"]:
        identifiers = []
        if ref.get("PMID"):
            identifiers.append(f"PMID `{ref['PMID']}`")
        if ref.get("DOI"):
            identifiers.append(f"DOI `{ref['DOI']}`")
        lines.append(f"| `{ref['id']}` | {', '.join(identifiers)} | {visible_bibliography_entry(ref)} |")

    lines.extend(["", "## Claim Mapping", "", "| Claim | Citation key | Support level | Notes |", "|---|---|---|---|"])
    for claim in project.get("claims", []):
        ref = refs.get(claim.get("support", ""), {})
        lines.append(
            f"| {claim.get('claim', '')} | `{ref.get('id', claim.get('support', ''))}` | "
            f"{claim.get('level', 'Not assessed')} | {claim.get('note', '')} |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_project(project: dict[str, Any], out_dir: Path) -> BuildOutputs:
    out_dir.mkdir(parents=True, exist_ok=True)
    outputs = output_paths(project, out_dir)
    make_placeholder_docx(project, outputs.placeholder_docx)
    make_active_docx(project, outputs.active_docx)
    make_csl_json(project, outputs.csl_json)
    make_ris(project, outputs.ris)
    make_claim_report(project, outputs.claim_report)
    return outputs


def inspect_docx_fields(docx: Path) -> dict[str, Any]:
    with zipfile.ZipFile(docx) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return {
        "citationFields": document_xml.count("ADDIN ZOTERO_ITEM"),
        "bibliographyFields": document_xml.count("ADDIN ZOTERO_BIBL"),
        "hasCSLCitationMarker": "CSL_CITATION" in document_xml,
        "hasCSLBibliographyMarker": "CSL_BIBLIOGRAPHY" in document_xml,
    }
