from __future__ import annotations

import json
import uuid
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "immunoglobulins_fibrosis_demo"
PLACEHOLDER_DOCX = OUT_DIR / "immunoglobulins_fibrosis_zotero_draft.docx"
EXPERIMENTAL_ACTIVE_DOCX = OUT_DIR / "immunoglobulins_fibrosis_zotero_active_field_experiment.docx"


REFERENCES = [
    {
        "key": "wynn-2012",
        "placeholder": "[ZOTERO: Wynn 2012]",
        "authors": [
            {"family": "Wynn", "given": "Thomas A."},
            {"family": "Ramalingam", "given": "Thirumalai R."},
        ],
        "ris_authors": ["Wynn, Thomas A.", "Ramalingam, Thirumalai R."],
        "title": "Mechanisms of fibrosis: therapeutic translation for fibrotic disease",
        "journal": "Nature Medicine",
        "short_journal": "Nat Med",
        "year": "2012",
        "date_parts": [2012, 7, 6],
        "volume": "18",
        "issue": "7",
        "pages": "1028-1040",
        "start_page": "1028",
        "end_page": "1040",
        "doi": "10.1038/nm.2807",
        "pmid": "22772564",
        "url": "https://doi.org/10.1038/nm.2807",
    },
    {
        "key": "stone-2012",
        "placeholder": "[ZOTERO: Stone 2012]",
        "authors": [
            {"family": "Stone", "given": "John H."},
            {"family": "Zen", "given": "Yoh"},
            {"family": "Deshpande", "given": "Vikram"},
        ],
        "ris_authors": ["Stone, John H.", "Zen, Yoh", "Deshpande, Vikram"],
        "title": "IgG4-related disease",
        "journal": "The New England Journal of Medicine",
        "short_journal": "N Engl J Med",
        "year": "2012",
        "date_parts": [2012, 2, 9],
        "volume": "366",
        "issue": "6",
        "pages": "539-551",
        "start_page": "539",
        "end_page": "551",
        "doi": "10.1056/NEJMra1104650",
        "pmid": "22316447",
        "url": "https://doi.org/10.1056/NEJMra1104650",
    },
    {
        "key": "deshpande-2012",
        "placeholder": "[ZOTERO: Deshpande 2012]",
        "authors": [
            {"family": "Deshpande", "given": "Vikram"},
            {"family": "Zen", "given": "Yoh"},
            {"family": "Chan", "given": "John KC"},
            {"family": "Yi", "given": "Eunhee E."},
            {"family": "Sato", "given": "Yasuharu"},
            {"family": "Yoshino", "given": "Tadashi"},
            {"family": "Kloppel", "given": "Gunter"},
            {"family": "Heathcote", "given": "J. Godfrey"},
            {"family": "Khosroshahi", "given": "Arezou"},
            {"family": "Ferry", "given": "Judith A."},
            {"family": "Aalberse", "given": "Rob C."},
            {"family": "Bloch", "given": "Donald B."},
            {"family": "Brugge", "given": "William R."},
            {"family": "Bateman", "given": "Adrian C."},
            {"family": "Carruthers", "given": "Mollie N."},
            {"family": "Chari", "given": "Suresh T."},
            {"family": "Cheuk", "given": "Wah"},
            {"family": "Cornell", "given": "Lynn D."},
            {"family": "Fernandez-Del Castillo", "given": "Carlos"},
            {"family": "Forcione", "given": "David G."},
            {"family": "Hamilos", "given": "Daniel L."},
            {"family": "Kamisawa", "given": "Terumi"},
            {"family": "Kasashima", "given": "Satomi"},
            {"family": "Kawa", "given": "Shigeyuki"},
            {"family": "Kawano", "given": "Mitsuhiro"},
            {"family": "Lauwers", "given": "Gregory Y."},
            {"family": "Masaki", "given": "Yasufumi"},
            {"family": "Nakanuma", "given": "Yasuni"},
            {"family": "Notohara", "given": "Kenji"},
            {"family": "Okazaki", "given": "Kazuichi"},
            {"family": "Ryu", "given": "Ji Kon"},
            {"family": "Saeki", "given": "Takako"},
            {"family": "Sahani", "given": "Dushyant V."},
            {"family": "Smyrk", "given": "Thomas C."},
            {"family": "Stone", "given": "James R."},
            {"family": "Takahira", "given": "Masayuki"},
            {"family": "Webster", "given": "George J."},
            {"family": "Yamamoto", "given": "Motohisa"},
            {"family": "Zamboni", "given": "Giuseppe"},
            {"family": "Umehara", "given": "Hisanori"},
            {"family": "Stone", "given": "John H."},
        ],
        "ris_authors": [
            "Deshpande, Vikram",
            "Zen, Yoh",
            "Chan, John KC",
            "Yi, Eunhee E.",
            "Sato, Yasuharu",
            "Yoshino, Tadashi",
            "Kloppel, Gunter",
            "Heathcote, J. Godfrey",
            "Khosroshahi, Arezou",
            "Ferry, Judith A.",
            "Aalberse, Rob C.",
            "Bloch, Donald B.",
            "Brugge, William R.",
            "Bateman, Adrian C.",
            "Carruthers, Mollie N.",
            "Chari, Suresh T.",
            "Cheuk, Wah",
            "Cornell, Lynn D.",
            "Fernandez-Del Castillo, Carlos",
            "Forcione, David G.",
            "Hamilos, Daniel L.",
            "Kamisawa, Terumi",
            "Kasashima, Satomi",
            "Kawa, Shigeyuki",
            "Kawano, Mitsuhiro",
            "Lauwers, Gregory Y.",
            "Masaki, Yasufumi",
            "Nakanuma, Yasuni",
            "Notohara, Kenji",
            "Okazaki, Kazuichi",
            "Ryu, Ji Kon",
            "Saeki, Takako",
            "Sahani, Dushyant V.",
            "Smyrk, Thomas C.",
            "Stone, James R.",
            "Takahira, Masayuki",
            "Webster, George J.",
            "Yamamoto, Motohisa",
            "Zamboni, Giuseppe",
            "Umehara, Hisanori",
            "Stone, John H.",
        ],
        "title": "Consensus statement on the pathology of IgG4-related disease",
        "journal": "Modern Pathology",
        "short_journal": "Mod Pathol",
        "year": "2012",
        "date_parts": [2012, 9],
        "volume": "25",
        "issue": "9",
        "pages": "1181-1192",
        "start_page": "1181",
        "end_page": "1192",
        "doi": "10.1038/modpathol.2012.72",
        "pmid": "22596100",
        "url": "https://doi.org/10.1038/modpathol.2012.72",
    },
]


CLAIMS = [
    {
        "claim": (
            "Fibrosis is framed as a persistent wound-healing program involving "
            "immune activation, fibroblast activation, and extracellular matrix remodeling."
        ),
        "support": "wynn-2012",
        "level": "Supported",
        "note": (
            "Review article on mechanisms of fibrosis and therapeutic translation; "
            "appropriate for broad fibrosis framing."
        ),
    },
    {
        "claim": (
            "Fibrotic tissue reflects ongoing crosstalk between inflammatory cells, "
            "stromal cells, and matrix-producing myofibroblasts."
        ),
        "support": "wynn-2012",
        "level": "Supported",
        "note": "Used as mechanistic background, not as evidence for an IgG-specific mechanism.",
    },
    {
        "claim": (
            "IgG4-related disease is a useful immunoglobulin-associated example of "
            "fibroinflammatory organ disease."
        ),
        "support": "stone-2012",
        "level": "Supported",
        "note": (
            "NEJM review describes IgG4-related disease as a fibroinflammatory condition "
            "with multi-organ involvement."
        ),
    },
    {
        "claim": (
            "The syndrome often includes dense lymphoplasmacytic infiltrates and "
            "increased IgG4-positive plasma cells."
        ),
        "support": "stone-2012",
        "level": "Supported",
        "note": "Kept general and tied to the disease overview source.",
    },
    {
        "claim": (
            "Diagnostic tissue interpretation emphasizes dense lymphoplasmacytic infiltrate, "
            "storiform fibrosis, obliterative phlebitis, and increased IgG4-positive plasma cells."
        ),
        "support": "deshpande-2012",
        "level": "Supported",
        "note": "Directly grounded in the pathology consensus statement.",
    },
]


def set_font(run, name: str = "Calibri", size: int | float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name: str, size: int | float, color: str | None = None, bold: bool | None = None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, "Calibri", 20, "0B2545", True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 11, "555555")
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, "2E74B5", True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, "2E74B5", True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)


def add_prose_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_font(run, size=11)
    return paragraph


def add_note_box(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.15
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    ppr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "2E74B5")
    border.append(left)
    ppr.append(border)
    label = paragraph.add_run("Zotero handoff: ")
    set_font(label, size=10.5, bold=True)
    body = paragraph.add_run(
        "Import the companion RIS or CSL JSON file into Zotero, replace each "
        "temporary placeholder with Zotero Word plugin Add/Edit Citation, then "
        "insert the bibliography under the References heading."
    )
    set_font(body, size=10.5)


def add_reference_insertion_hint(doc: Document):
    heading = doc.add_heading("References", level=1)
    heading.paragraph_format.keep_with_next = True
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(
        "[After replacing all placeholders with active Zotero citations, place the "
        "cursor below this paragraph and click Zotero > Add/Edit Bibliography.]"
    )
    set_font(run, size=10.5)
    run.italic = True


def csl_item_data(ref: dict) -> dict:
    return {
        "id": ref["key"],
        "type": "article-journal",
        "title": ref["title"],
        "container-title": ref["journal"],
        "container-title-short": ref["short_journal"],
        "author": ref["authors"],
        "issued": {"date-parts": [ref["date_parts"]]},
        "volume": ref["volume"],
        "issue": ref["issue"],
        "page": ref["pages"],
        "DOI": ref["doi"],
        "PMID": ref["pmid"],
        "URL": ref["url"],
    }


def import_log_item_keys_by_doi() -> dict[str, str]:
    path = OUT_DIR / "zotero_import_result.json"
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}
    return {
        item["DOI"].lower(): item["key"]
        for item in data.get("items", [])
        if item.get("DOI") and item.get("key")
    }


def zotero_like_uris(ref: dict, keys_by_doi: dict[str, str]) -> list[str]:
    item_key = keys_by_doi.get(ref["doi"].lower(), ref["key"].replace("-", "").upper()[:8])
    return [f"http://zotero.org/users/local/citationtool/items/{item_key}"]


def zotero_citation_json(ref: dict, number: int, keys_by_doi: dict[str, str]) -> str:
    visible = f"[{number}]"
    payload = {
        "citationID": f"citationtool-{uuid.uuid5(uuid.NAMESPACE_URL, ref['doi']).hex[:12]}",
        "citationItems": [
            {
                "id": ref["key"],
                "uris": zotero_like_uris(ref, keys_by_doi),
                "itemData": csl_item_data(ref),
            }
        ],
        "properties": {
            "formattedCitation": visible,
            "plainCitation": visible,
            "noteIndex": 0,
        },
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return json.dumps(payload, separators=(",", ":"), ensure_ascii=True)


def append_text_element(run_element, text: str):
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)


def append_field_run(paragraph, instruction: str, result_text: str):
    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    paragraph._p.append(begin)

    for idx in range(0, len(instruction), 900):
        instr_run = OxmlElement("w:r")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction[idx : idx + 900]
        instr_run.append(instr_text)
        paragraph._p.append(instr_run)

    separate = OxmlElement("w:r")
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate.append(fld_separate)
    paragraph._p.append(separate)

    for line_index, line in enumerate(result_text.split("\n")):
        result_run = OxmlElement("w:r")
        if line_index:
            br = OxmlElement("w:br")
            result_run.append(br)
        append_text_element(result_run, line)
        paragraph._p.append(result_run)

    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._p.append(end)


def append_zotero_citation(paragraph, ref: dict, number: int, keys_by_doi: dict[str, str]):
    instruction = f" ADDIN ZOTERO_ITEM CSL_CITATION {zotero_citation_json(ref, number, keys_by_doi)}"
    append_field_run(paragraph, instruction, f"[{number}]")


def append_zotero_bibliography(paragraph, bibliography_text: str):
    code = json.dumps({"uncited": [], "omitted": [], "custom": []}, separators=(",", ":"))
    instruction = f" ADDIN ZOTERO_BIBL {code} CSL_BIBLIOGRAPHY"
    append_field_run(paragraph, instruction, bibliography_text)


def add_active_prose_paragraph(doc: Document, parts: list[str | tuple[str, int]], keys_by_doi: dict[str, str]):
    refs_by_key = {ref["key"]: ref for ref in REFERENCES}
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.25
    for part in parts:
        if isinstance(part, tuple):
            ref_key, number = part
            append_zotero_citation(paragraph, refs_by_key[ref_key], number, keys_by_doi)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=11)
    return paragraph


def make_docx():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Immunoglobulins and Fibrosis: Zotero-Ready Draft")
    set_font(title_run, size=20, bold=True)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "Prototype text with verified sources, temporary Zotero placeholders, "
        "and companion import files"
    )
    set_font(subtitle_run, size=11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("Draft Introduction", level=1)
    add_prose_paragraph(
        doc,
        "Fibrosis is best understood as a persistent wound-healing program in "
        "which immune activation, fibroblast activation, and extracellular matrix "
        "remodeling fail to resolve [ZOTERO: Wynn 2012]. Rather than being a "
        "passive scar, fibrotic tissue reflects ongoing crosstalk between "
        "inflammatory cells, tissue-resident stromal cells, and matrix-producing "
        "myofibroblasts [ZOTERO: Wynn 2012].",
    )
    add_prose_paragraph(
        doc,
        "Within this broader biology, IgG4-related disease provides a useful "
        "example of how an immunoglobulin-associated immune response can travel "
        "with organ fibrosis. The syndrome is characterized by tumor-like "
        "fibroinflammatory lesions across multiple organs, often with dense "
        "lymphoplasmacytic infiltrates and increased IgG4-positive plasma cells "
        "[ZOTERO: Stone 2012].",
    )
    add_prose_paragraph(
        doc,
        "For a citation workflow, the most important distinction is that pathology "
        "claims should be anchored to tissue-based criteria rather than to serum "
        "IgG4 alone. Consensus pathology criteria emphasize dense lymphoplasmacytic "
        "infiltrate, storiform fibrosis, obliterative phlebitis, and increased "
        "IgG4-positive plasma cells in the appropriate clinical context "
        "[ZOTERO: Deshpande 2012].",
    )

    add_note_box(doc)
    add_reference_insertion_hint(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Generated demo draft - replace placeholders with Zotero fields")
    set_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor.from_string("555555")

    doc.save(PLACEHOLDER_DOCX)


def make_experimental_active_docx():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    keys_by_doi = import_log_item_keys_by_doi()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Immunoglobulins and Fibrosis: Experimental Zotero Field Draft")
    set_font(title_run, size=20, bold=True)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "Generated Word fields using Zotero-style ADDIN codes; verify once in Word with Zotero Refresh"
    )
    set_font(subtitle_run, size=11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("Draft Introduction", level=1)
    add_active_prose_paragraph(
        doc,
        [
            "Fibrosis is best understood as a persistent wound-healing program in which immune activation, fibroblast activation, and extracellular matrix remodeling fail to resolve ",
            ("wynn-2012", 1),
            ". Rather than being a passive scar, fibrotic tissue reflects ongoing crosstalk between inflammatory cells, tissue-resident stromal cells, and matrix-producing myofibroblasts ",
            ("wynn-2012", 1),
            ".",
        ],
        keys_by_doi,
    )
    add_active_prose_paragraph(
        doc,
        [
            "Within this broader biology, IgG4-related disease provides a useful example of how an immunoglobulin-associated immune response can travel with organ fibrosis. The syndrome is characterized by tumor-like fibroinflammatory lesions across multiple organs, often with dense lymphoplasmacytic infiltrates and increased IgG4-positive plasma cells ",
            ("stone-2012", 2),
            ".",
        ],
        keys_by_doi,
    )
    add_active_prose_paragraph(
        doc,
        [
            "For a citation workflow, the most important distinction is that pathology claims should be anchored to tissue-based criteria rather than to serum IgG4 alone. Consensus pathology criteria emphasize dense lymphoplasmacytic infiltrate, storiform fibrosis, obliterative phlebitis, and increased IgG4-positive plasma cells in the appropriate clinical context ",
            ("deshpande-2012", 3),
            ".",
        ],
        keys_by_doi,
    )

    doc.add_heading("References", level=1)
    bibliography_text = "\n".join(f"{index}. {format_ref(ref)}" for index, ref in enumerate(REFERENCES, 1))
    bibliography = doc.add_paragraph()
    bibliography.paragraph_format.space_after = Pt(6)
    bibliography.paragraph_format.line_spacing = 1.15
    append_zotero_bibliography(bibliography, bibliography_text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Experimental generated Zotero fields - validate with Zotero Refresh")
    set_font(footer_run, size=9)
    footer_run.font.color.rgb = RGBColor.from_string("555555")

    doc.save(EXPERIMENTAL_ACTIVE_DOCX)


def make_csl_json():
    entries = []
    for ref in REFERENCES:
        entries.append(csl_item_data(ref))
    path = OUT_DIR / "immunoglobulins_fibrosis_zotero_references.csl.json"
    path.write_text(json.dumps(entries, indent=2) + "\n", encoding="utf-8")


def make_ris():
    blocks: list[str] = []
    for ref in REFERENCES:
        lines = ["TY  - JOUR", f"ID  - {ref['key']}"]
        for author in ref["ris_authors"]:
            lines.append(f"AU  - {author}")
        lines.extend(
            [
                f"TI  - {ref['title']}",
                f"T2  - {ref['journal']}",
                f"JO  - {ref['short_journal']}",
                f"PY  - {ref['year']}",
                f"VL  - {ref['volume']}",
                f"IS  - {ref['issue']}",
                f"SP  - {ref['start_page']}",
                f"EP  - {ref['end_page']}",
                f"DO  - {ref['doi']}",
                f"UR  - {ref['url']}",
                f"AN  - PMID:{ref['pmid']}",
            ]
        )
        lines.append("ER  -")
        blocks.append("\n".join(lines))
    path = OUT_DIR / "immunoglobulins_fibrosis_zotero_references.ris"
    path.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")


def format_ref(ref):
    first = ref["ris_authors"][0].split(",")[0]
    author_label = f"{first} et al." if len(ref["ris_authors"]) > 2 else " and ".join(
        name.split(",")[0] for name in ref["ris_authors"]
    )
    return (
        f"{author_label} {ref['year']}. {ref['title']}. "
        f"{ref['short_journal']} {ref['volume']}({ref['issue']}):{ref['pages']}. "
        f"doi:{ref['doi']}; PMID:{ref['pmid']}."
    )


def repo_path(path: Path) -> str:
    return str(path.relative_to(ROOT))


def make_claim_report():
    refs_by_key = {ref["key"]: ref for ref in REFERENCES}
    lines = [
        "# Claim Support Report",
        "",
        "Generated for the Zotero-ready immunoglobulins/fibrosis demo.",
        "",
        "## Verification Summary",
        "",
        "All three references were selected because PubMed metadata confirms the PMID, DOI, journal, year, volume, issue, and pages. The Nature Medicine DOI was also checked through Crossref during generation setup.",
        "",
        "| Key | Zotero placeholder | Verified identifier | Reference |",
        "|---|---|---|---|",
    ]
    for ref in REFERENCES:
        lines.append(
            f"| `{ref['key']}` | `{ref['placeholder']}` | PMID `{ref['pmid']}`, DOI `{ref['doi']}` | {format_ref(ref)} |"
        )

    lines.extend(
        [
            "",
            "## Claim Mapping",
            "",
            "| Claim | Placeholder | Support level | Notes |",
            "|---|---|---|---|",
        ]
    )
    for item in CLAIMS:
        ref = refs_by_key[item["support"]]
        lines.append(
            f"| {item['claim']} | `{ref['placeholder']}` | {item['level']} | {item['note']} |"
        )

    lines.extend(
        [
            "",
            "## Zotero Round Trip",
            "",
            "Recommended automation:",
            "",
            "```bash",
            "python3 tools/run_zotero_ready_demo.py --open-active-word",
            "```",
            "",
            "That command regenerates both Word drafts, imports the references into the currently selected Zotero target via Zotero's local connector, and opens the experimental active-field draft. In Word, run Zotero `Refresh`; if Zotero can edit the citations and bibliography, this becomes the automated path.",
            "",
            "Conservative placeholder workflow:",
            "",
            "```bash",
            "python3 tools/run_zotero_ready_demo.py --open-word",
            "```",
            "",
            "This opens the safe placeholder draft. The remaining manual step is to use the Zotero Word plugin to replace the placeholders with active citation fields.",
            "",
            "Experimental automation output:",
            "",
            f"- `{repo_path(EXPERIMENTAL_ACTIVE_DOCX)}` contains generated Word ADDIN field codes using Zotero's `ZOTERO_ITEM` and `ZOTERO_BIBL` markers. Open it in Word and run Zotero `Refresh`; if the citations open in Zotero's edit dialog, this can become the automated path.",
            f"- `python3 tools/run_zotero_ready_demo.py --refresh-active-word` asks Zotero's Mac Word integration endpoint to run `Refresh` on `{repo_path(EXPERIMENTAL_ACTIVE_DOCX)}`.",
            "",
            "Manual fallback:",
            "",
            "1. Import `immunoglobulins_fibrosis_zotero_references.ris` or `immunoglobulins_fibrosis_zotero_references.csl.json` into Zotero.",
            "2. Open `immunoglobulins_fibrosis_zotero_draft.docx` in Microsoft Word.",
            "3. Replace each `[ZOTERO: ...]` placeholder with Zotero Word plugin `Add/Edit Citation`.",
            "4. Under the References heading, use Zotero Word plugin `Add/Edit Bibliography`.",
            "5. Save the result as a new Word file; that copy will contain active Zotero fields.",
            "",
            "Important: the generated draft intentionally does not contain handcrafted Zotero field codes. The placeholders are plain text until the Zotero Word plugin replaces them.",
            "",
        ]
    )
    (OUT_DIR / "claim_support_report.md").write_text("\n".join(lines), encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_docx()
    make_experimental_active_docx()
    make_ris()
    make_csl_json()
    make_claim_report()


if __name__ == "__main__":
    main()
